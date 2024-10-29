import datetime
import requests
import re
import zipfile
import pandas as pd
import os
import fitz
import json
from bs4 import BeautifulSoup
from docx import Document

class ProposalSkripsiByGemini:
    def __init__(self, judul_skripsi, sejak_tahun=None, api_key="", hasil_file="skripsi.docx"):
        self.judul_skripsi = judul_skripsi
        self.sejak_tahun = sejak_tahun
        if not self.sejak_tahun:
            self.sejak_tahun = datetime.datetime.now().year - 5
        else:
            self.sejak_tahun = datetime.datetime.now().year if not self.sejak_tahun else self.sejak_tahun
        self.api_key = api_key
        self.hasil_file = hasil_file
        self.prompt = {
            "roleplay": "Kamu adalah seorang mahasiswa tingkat akhir yang mahir yang sering melakukan banyak penelitian",
            "ingat": "Berikut adalah paper yang akan dibahas. selanjutnya, tolong review namun tidak perlu di respon lebih jauh",
            "pendahuluan": "Buatkan pendahuluan dalam bentuk paragraf dan memiliki format sitasi APA dengan berdasarkan paper sebelumnya dengan judul ",
            "sitasi": "Buatkan sitasi yang mengarah ke paper ini dengan format APA:",
            "masalah": "Buatkan research problem berdasarkan mengarah paper-paper sebelumnya dan dengan judul ",
            "tujuan": "Buatkan tujuan penelitian berdasarkan mengarah paper-paper sebelumnya dan dengan judul",
            "batasan": "Buatkan batasan penelitian berdasarkan mengarah paper-paper sebelumnya dan dengan judul",
            "keyword": "Buatkan keyword-keyword yang berkaitan dengan judul ",
            "studi-pustaka1": "Buatkan studi pustaka berdasarkan paper-paper yang sudah direview sebelumnya",
            "studi-pustaka2": "Buatkan studi pustaka berdasarkan keyword-keyword sebelumnya",
            "metodpen": "Buatkan metodologi penelitian berdasarkan mengarah paper-paper sebelumnya dan dengan judul",
            "daftar-pustaka" : "Buatkan daftar pustaka dengan format APA berdasarkan paper-paper sebelumnya"
        }

    def scrape_links_paper(self, url):
        response = requests.get(url)
        soup = BeautifulSoup(response.content, "html.parser")
        h3_elements = soup.find_all("h3", class_="gs_rt")
        links = [h3_element.find("a")["href"] for h3_element in h3_elements if h3_element.find("a")]
        return links

    def scrape_all_page(self, url, start, end):
        return [url + "&start=" + str(i * 10) for i in range(start, end + 1)]

    def flatten_list(self, nested_list):
        flattened_list = []
        for element in nested_list:
            if isinstance(element, list):
                flattened_list.extend(self.flatten_list(element))
            else:
                flattened_list.append(element)
        return flattened_list

    def remove_symbols(self, text):
        symbol_regex = re.compile(r"[^\w\s]")
        return symbol_regex.sub("", text)

    def create_zip_archive(self, files, filename):
        with zipfile.ZipFile(filename, 'w') as zip_file:
            for file in files:
                zip_file.write(file)

    def get_all_links(self, url, start, end, pdf_dir="pdf_dir"):
        all_res_file = []
        all_url_page = self.scrape_all_page(url, start, end)
        all_links = []
        for url_ in all_url_page:
            all_links.extend(self.scrape_links_paper(url_))
        all_links = self.flatten_list(all_links)
        df = pd.DataFrame({"id": list(range(1, len(all_links) + 1)), "url": all_links})
        filename = self.remove_symbols(url)
        df.to_csv(f"{filename}.csv", index=False)
        df.to_excel(f"{filename}.xlsx", index=False)
        all_res_file += [f"{filename}.csv", f"{filename}.xlsx"]
        all_link_of_pdf = [url_ for url_ in all_links if ".pdf" in url_]
        df_pdf = pd.DataFrame({"id": list(range(1, len(all_link_of_pdf) + 1)), "url": all_link_of_pdf})
        df_pdf.to_csv(f"{filename}_pdf.csv", index=False)
        df_pdf.to_excel(f"{filename}_pdf.xlsx", index=False)
        all_res_file += [f"{filename}_pdf.csv", f"{filename}_pdf.xlsx"]
        self.create_zip_archive(all_res_file, f"{filename}.zip")
        for file in all_res_file:
            try:
                os.remove(file)
            except OSError:
                pass
        os.makedirs(pdf_dir, exist_ok=True)
        for link in all_link_of_pdf:
            try:
                response = requests.get(link)
                with open(os.path.join(pdf_dir, os.path.basename(link)), 'wb') as f:
                    f.write(response.content)
            except:
                pass
        pdf_files = [os.path.join(pdf_dir, file_) for file_ in os.listdir(pdf_dir) if file_.endswith('.pdf')]
        self.create_zip_archive([os.path.join(pdf_dir, file_) for file_ in os.listdir(pdf_dir)], f"{filename}_pdf.zip")
        return pdf_files, f"{filename}.zip", f"{filename}_pdf.zip"

    def extract_text_from_pdf(self, pdf_path):
        doc = fitz.open(pdf_path)
        text = "".join([page.get_text() for page in doc])
        return text

    def _gemini(self, prompt, data):
        url = f'https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent?key={self.api_key}'
        headers = {'Content-Type': 'application/json'}
        prompt = {"parts": [{"text": prompt}], 'role': 'user'}
        data['contents'].append(prompt)
        response = requests.post(url, headers=headers, json=data)
        if response.status_code == 200:
            result = response.json()
            data['contents'].append(result['candidates'][0]['content'])
            return data
        else:
            return data

    def export_to_docx(self, title, text, filename="skripsi.docx"):
        doc = Document()
        doc.add_heading(title, 0)
        lines = text.splitlines()
        for line in lines:
            if line.startswith("##"):
                doc.add_heading(line[2:].strip(), level=1).bold = True
            elif line.startswith("**") and line.endswith("**"):
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(line[2:-2].strip())
                run.bold = True
            elif line.startswith("* "):
                paragraph = doc.add_paragraph(style='List Bullet')
                paragraph.add_run(line[2:].strip())
            else:
                doc.add_paragraph(line)
        doc.save(filename)

    def run(self):
        q_scholar = self.judul_skripsi.replace(" ", "+")
        URL = f"https://scholar.google.com/scholar?as_ylo={self.sejak_tahun}&q={q_scholar}"
        pdf_dir = self.judul_skripsi + "_pdf_dir"
        all_pdf_files, links_zip, pdf_link_zip = self.get_all_links(URL, 0, 20, pdf_dir)
        all_text = []
        for pdf_file in all_pdf_files:
            try:
              text = self.extract_text_from_pdf(pdf_file)
              all_text.append(text)
            except:
              pass
        review_paper = {"contents": []}
        for text in all_text:
            prompt_ = self.prompt['roleplay'] + "\n\n" + self.prompt['ingat'] + "\n```" + text + "```"
            review_paper = self._gemini(prompt_, review_paper)
        with open('review_paper.json', 'w') as f:
            json.dump(review_paper, f)
        pendahuluan = review_paper
        prompt_ = self.prompt['roleplay'] + "\n\n" + self.prompt['daftar-pustaka']
        pendahuluan = self._gemini(prompt_, pendahuluan)
        daftar_pustaka = pendahuluan['contents'][-1]['parts'][0]['text']
        prompt_ = self.prompt['roleplay'] + "\n\n" + self.prompt['pendahuluan'] + self.judul_skripsi
        pendahuluan = self._gemini(prompt_, pendahuluan)
        pendahuluan_skripsi = pendahuluan['contents'][-1]['parts'][0]['text']
        prompt_ = self.prompt['roleplay'] + "\n\n" + self.prompt['masalah'] + self.judul_skripsi
        pendahuluan = self._gemini(prompt_, pendahuluan)
        masalah_skripsi = pendahuluan['contents'][-1]['parts'][0]['text']
        prompt_ = self.prompt['roleplay'] + "\n\n" + self.prompt['tujuan'] + self.judul_skripsi
        pendahuluan = self._gemini(prompt_, pendahuluan)
        tujuan_skripsi = pendahuluan['contents'][-1]['parts'][0]['text']
        prompt_ = self.prompt['roleplay'] + "\n\n" + self.prompt['batasan'] + self.judul_skripsi
        pendahuluan = self._gemini(prompt_, pendahuluan)
        batasan_skripsi = pendahuluan['contents'][-1]['parts'][0]['text']
        with open('pendahuluan.json', 'w') as f:
            json.dump(pendahuluan, f)
        studi_pustaka = pendahuluan
        prompt_ = self.prompt['roleplay'] + "\n\n" + self.prompt['studi-pustaka1'] + self.judul_skripsi
        studi_pustaka = self._gemini(prompt_, studi_pustaka)
        studi_pustaka1 = studi_pustaka['contents'][-1]['parts'][0]['text']
        prompt_ = self.prompt['roleplay'] + "\n\n" + self.prompt['keyword'] + self.judul_skripsi
        studi_pustaka = self._gemini(prompt_, studi_pustaka)
        keyword = studi_pustaka['contents'][-1]['parts'][0]['text']
        prompt_ = self.prompt['roleplay'] + "\n\n" + self.prompt['studi-pustaka2']
        studi_pustaka = self._gemini(prompt_, studi_pustaka)
        studi_pustaka2 = studi_pustaka['contents'][-1]['parts'][0]['text']
        with open('studi_pustaka.json', 'w') as f:
            json.dump(studi_pustaka, f)
        metodpen = studi_pustaka
        prompt_ = self.prompt['roleplay'] + "\n\n" + self.prompt['metodpen'] + self.judul_skripsi
        metodpen = self._gemini(prompt_, metodpen)
        metodologi_penelitian = metodpen['contents'][-1]['parts'][0]['text']
        skripsi = f"""
Judul Skripsi : {self.judul_skripsi}
\n\n
## BAB I
## Pendahuluan
\n\n
## Latar Belakang
{pendahuluan_skripsi}

## Masalah Penelitian
{masalah_skripsi}

## Tujuan Penelitian
{tujuan_skripsi}

## Batasan Penelitian
{batasan_skripsi}
\n\n
## BAB II
## Penelitian Sebelumnya
{studi_pustaka1}

## Keyword
{keyword}

## Studi Pustaka
{studi_pustaka2}
\n\n
## BAB III
## Metodologi Penelitian

## Metodologi Penelitian
{metodologi_penelitian}
\n\n
## Daftar Pustaka
{daftar_pustaka}
"""
        self.export_to_docx(self.judul_skripsi, skripsi, self.hasil_file)
        return self.hasil_file, links_zip, pdf_link_zip
        
